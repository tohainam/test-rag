"""
File Content Extractor Service

Extracts text content from various file formats (PDF, DOCX, TXT, MD).
"""

import io
import logging
from typing import Tuple, Dict
import pdfplumber
from docx import Document

from src.services.minio_service import MinIOService

logger = logging.getLogger(__name__)


class FileContentExtractor:
    """Service for extracting text content from files stored in MinIO"""

    def __init__(self, minio_service: MinIOService):
        """
        Initialize the file content extractor.

        Args:
            minio_service: MinIO service instance for file access
        """
        self.minio_service = minio_service

    def extract_content(
        self,
        filename: str,
        content_type: str,
        minio_bucket: str,
        minio_object_name: str
    ) -> Tuple[str, Dict]:
        """
        Extract text content from a file based on its content type.

        Args:
            filename: Original filename
            content_type: MIME type of the file
            minio_bucket: MinIO bucket name
            minio_object_name: Object name in MinIO

        Returns:
            Tuple of (extracted_text, metadata_dict)
            metadata includes: char_count, word_count, page_count (for PDFs)

        Raises:
            ValueError: If content type is not supported
            Exception: If extraction fails
        """
        logger.info(f"Extracting content from file: {filename} (type: {content_type})")

        try:
            # Download file from MinIO
            file_data = self._download_file(minio_bucket, minio_object_name)

            # Extract based on content type
            if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
                text, metadata = self._extract_from_pdf(file_data)
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword"
            ] or filename.lower().endswith((".docx", ".doc")):
                text, metadata = self._extract_from_docx(file_data)
            elif content_type == "text/plain" or filename.lower().endswith(".txt"):
                text, metadata = self._extract_from_txt(file_data)
            elif content_type == "text/markdown" or filename.lower().endswith(".md"):
                text, metadata = self._extract_from_md(file_data)
            else:
                raise ValueError(f"Unsupported content type: {content_type}")

            # Add common metadata
            metadata["filename"] = filename
            metadata["content_type"] = content_type

            logger.info(
                f"Successfully extracted {metadata['char_count']} characters "
                f"({metadata['word_count']} words) from {filename}"
            )

            return text, metadata

        except Exception as e:
            logger.error(f"Failed to extract content from {filename}: {str(e)}")
            raise

    def _download_file(self, bucket: str, object_name: str) -> bytes:
        """Download file from MinIO as bytes"""
        try:
            response = self.minio_service.client.get_object(bucket, object_name)
            file_data = response.read()
            response.close()
            response.release_conn()
            return file_data
        except Exception as e:
            logger.error(f"Failed to download file from MinIO: {str(e)}")
            raise

    def _extract_from_pdf(self, file_data: bytes) -> Tuple[str, Dict]:
        """Extract text from PDF file using pdfplumber"""
        try:
            file_obj = io.BytesIO(file_data)
            text_parts = []
            page_count = 0

            with pdfplumber.open(file_obj) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            full_text = "\n\n".join(text_parts)
            metadata = {
                "char_count": len(full_text),
                "word_count": len(full_text.split()),
                "page_count": page_count,
            }

            return full_text, metadata

        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    def _extract_from_docx(self, file_data: bytes) -> Tuple[str, Dict]:
        """Extract text from DOCX file using python-docx"""
        try:
            file_obj = io.BytesIO(file_data)
            doc = Document(file_obj)

            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            full_text = "\n\n".join(text_parts)
            metadata = {
                "char_count": len(full_text),
                "word_count": len(full_text.split()),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }

            return full_text, metadata

        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    def _extract_from_txt(self, file_data: bytes) -> Tuple[str, Dict]:
        """Extract text from plain text file"""
        try:
            # Try UTF-8 first, fallback to latin-1
            try:
                text = file_data.decode("utf-8")
            except UnicodeDecodeError:
                text = file_data.decode("latin-1")

            metadata = {
                "char_count": len(text),
                "word_count": len(text.split()),
                "line_count": len(text.splitlines()),
            }

            return text, metadata

        except Exception as e:
            logger.error(f"TXT extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from TXT: {str(e)}")

    def _extract_from_md(self, file_data: bytes) -> Tuple[str, Dict]:
        """Extract text from Markdown file (same as txt)"""
        try:
            # Markdown files are plain text
            try:
                text = file_data.decode("utf-8")
            except UnicodeDecodeError:
                text = file_data.decode("latin-1")

            metadata = {
                "char_count": len(text),
                "word_count": len(text.split()),
                "line_count": len(text.splitlines()),
            }

            return text, metadata

        except Exception as e:
            logger.error(f"MD extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from MD: {str(e)}")

    def estimate_question_count(self, text: str) -> int:
        """
        Estimate appropriate number of questions based on content length.

        Args:
            text: Extracted text content

        Returns:
            Recommended number of questions to generate
        """
        word_count = len(text.split())

        # Heuristic: roughly 1 question per 200-300 words
        # Minimum 3, maximum 15 questions per file
        if word_count < 500:
            return 3
        elif word_count < 1500:
            return 5
        elif word_count < 3000:
            return 8
        elif word_count < 5000:
            return 10
        else:
            return min(15, max(3, word_count // 300))
